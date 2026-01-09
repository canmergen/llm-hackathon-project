# =============================================================================
# Banka Dokümanları Preparation Script - Word Only
# =============================================================================
"""
Banka Dokümanları Preparation Script - Word Only

This script processes Word documents (.docx) to extract,
chunk, and prepare vector-ready data for AI applications.

Features:
- Word document parsing with paragraph tracking
- Text chunking for retrieval systems
- Vector-ready subchunking with embedding optimization
- Comprehensive quality control and logging

Author: Can Mergen
Date: 2026-01-02
Version: 1.0.0
"""

# =============================================================================
# IMPORTS
# =============================================================================

import sys
import logging
from pathlib import Path
import re
import json
from typing import List, Dict, Any, Optional
from collections import Counter

# Third-party imports
try:
    from docx import Document
except ImportError:
    Document = None

# Local imports
from scope_utils import assign_bank_document_scopes

# =============================================================================
# CONSTANTS
# =============================================================================

PROJECT_MARKERS = {"data", "src"}

# Chunking parameters
MAX_EMBEDDING_CHARS = 900
MIN_SUBCHUNK_CHARS = 80

# File paths (will be set after directory detection)
BANKA_DOKUMANLARI_RAW_DIR: Optional[Path] = None
BANKA_DOKUMANLARI_PARSED_DIR: Optional[Path] = None
BANKA_DOKUMANLARI_CHUNKS_DIR: Optional[Path] = None

# =============================================================================
# REGEX PATTERNS
# =============================================================================

# Section patterns - A., B., C. or I., II., III. or 1., 2., 3.
MAIN_SECTION_RE = re.compile(r"^([A-Z])\.([\w\sÇĞİÖŞÜçğıöşü]+)$")
ROMAN_SECTION_RE = re.compile(r"^([IVX]+)\.([\w\sÇĞİÖŞÜçğıöşü]+)$")

# Article patterns (MADDE, Madde)
MADDE_RE = re.compile(r"^(?:Madde|MADDE)\s+(\d+)[:\-–]?\s*(.*)$")

# Numbered/lettered item patterns
NUMBERED_ITEM_RE = re.compile(r"^(\d+)\.\s+(.+)$")
LETTERED_ITEM_RE = re.compile(r"^([a-zçğıöşü])\)\s+(.+)$", re.IGNORECASE)

# Noise patterns
NOISE_TOKEN_RE = re.compile(
    r"(\b\d{8,}\b|"            # very long numeric strings
    r"\b\d{2}:\d{2}:\d{2}\b|" # timestamps
    r"Sayfa\s+\d+|"            # Page numbers
    r"^\d+/\d+$)"              # Page numbers like "1/9"
    , re.IGNORECASE
)

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def find_project_root(start_path: Optional[Path] = None) -> Path:
    """Find the project root directory by searching for marker directories."""
    start_path = (start_path or Path.cwd()).resolve()
    for path in [start_path, *start_path.parents]:
        if PROJECT_MARKERS.issubset({item.name for item in path.iterdir() if item.is_dir()}):
            return path
    raise RuntimeError(f"Project root not found from: {start_path}")

def setup_logging(log_file_path: Path) -> logging.Logger:
    """Configure logging with both file and console handlers."""
    log_file_path.parent.mkdir(exist_ok=True)
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file_path, mode='w', encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    
    return logging.getLogger(__name__)

def setup_directories(base_dir: Path) -> None:
    """Set up global directory paths."""
    global BANKA_DOKUMANLARI_RAW_DIR, BANKA_DOKUMANLARI_PARSED_DIR, BANKA_DOKUMANLARI_CHUNKS_DIR
    
    banka_dokumanlari_dir = base_dir / "data" / "banka_dokumanlari"
    BANKA_DOKUMANLARI_RAW_DIR = banka_dokumanlari_dir / "raw"
    BANKA_DOKUMANLARI_PARSED_DIR = banka_dokumanlari_dir / "parsed"
    BANKA_DOKUMANLARI_CHUNKS_DIR = banka_dokumanlari_dir / "chunks"
    
    # Create directories if they don't exist
    BANKA_DOKUMANLARI_PARSED_DIR.mkdir(parents=True, exist_ok=True)
    BANKA_DOKUMANLARI_CHUNKS_DIR.mkdir(parents=True, exist_ok=True)

def safe_text(text: str) -> str:
    """Clean and normalize text for processing."""
    if not text:
        return ""
    text = text.replace("\r", "")
    return re.sub(r"\n{3,}", "\n\n", text).strip()

def normalize_for_embedding(text: str) -> str:
    """Normalize text for embedding by converting newlines to spaces."""
    if not text:
        return ""
    text = text.replace("\r", "")
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()

def clean_noise_tokens(text: str) -> str:
    """Remove known noise tokens."""
    if not text:
        return ""
    cleaned = NOISE_TOKEN_RE.sub("", text)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip()

def determine_doc_type(doc_id: str) -> str:
    """Determine document type from doc_id."""
    doc_id_lower = doc_id.lower()
    if 'sozlesme' in doc_id_lower or 'sozlesmesi' in doc_id_lower:
        return 'contract'
    elif 'kredi' in doc_id_lower and 'kart' in doc_id_lower:
        return 'credit_card'
    elif 'ucret' in doc_id_lower or 'tarife' in doc_id_lower:
        return 'fee_schedule'
    elif 'politika' in doc_id_lower:
        return 'policy'
    return 'word_document'

# =============================================================================
# EXCEL ATTACHMENT HELPERS
# =============================================================================

def _excel_token_candidates(stem: str) -> List[str]:
    """Generate simple tokens from excel filename stem for matching in Word text."""
    stem_lower = stem.lower()
    token_space = stem.replace("_", " ").lower()
    no_digits_space = re.sub(r"\d+", "", token_space).strip()
    no_digits_underscore = re.sub(r"\d+", "", stem_lower).strip()
    return [t for t in {stem_lower, token_space, no_digits_space, no_digits_underscore} if t]

def find_referenced_excel_records(word_full_text: str, base_dir: Path, current_doc_id: str) -> List[Dict[str, Any]]:
    """Return excel vector records whose filename is mentioned in Word text."""
    chunks_dir = base_dir / "data" / "banka_dokumanlari" / "chunks"
    word_text_l = word_full_text.lower()
    records: List[Dict[str, Any]] = []

    for path in chunks_dir.glob("*_vector_ready.json"):
        if current_doc_id in path.stem:
            continue  # skip self
        stem = path.stem.replace("_vector_ready", "")
        tokens = _excel_token_candidates(stem)
        if not any(tok and tok in word_text_l for tok in tokens):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Tag source for clarity
            for rec in data:
                rec = dict(rec)
                rec.setdefault("metadata", {})
                rec["metadata"]["source"] = "excel_attachment"
                records.append(rec)
            logger.info(f"Attached excel vector records from {path.name} (match tokens: {tokens})")
        except Exception as exc:
            logger.warning(f"Could not load excel vector file {path}: {exc}")
            continue
    return records

# =============================================================================
# SCOPE MAPPING
# =============================================================================

def guess_scope(section_name: Optional[str], doc_id: str = "") -> List[str]:
    """Map section/heading names to teblig benzeri scope etiketleri."""
    scopes = set()
    
    # Doküman ID'den scope çıkar (kredi kartı sözleşmesi ise)
    if doc_id and "kredi" in doc_id.lower() and "kart" in doc_id.lower():
        scopes.add("credit_card")
    
    if not section_name:
        return sorted(scopes) if scopes else ["general_contract"]
    
    name = section_name.lower()

    # Para transferleri / EFT / FAST / IBAN
    if any(k in name for k in ["eft", "fast", "havale", "para transfer", "swift", "iban"]):
        scopes.add("money_transfer")
        scopes.add("eft_havale_fast")

    # Mevduat hesapları
    if any(k in name for k in ["vadesiz", "vadeli", "mevduat", "tevd", "döviz", "turuncu", "e-turuncu", "e turuncu"]):
        scopes.add("deposit_account")

    # Altın / kıymetli maden
    if any(k in name for k in ["altın", "altin", "kıymetli", "kiymetli"]):
        scopes.add("precious_metal_transfer")

    # Kredi kartı
    if any(k in name for k in ["kredi kart", "banka kart", "kart"]):
        scopes.add("credit_card")
        scopes.add("credit_card")
    if any(k in name for k in ["nakit avans", "nakit çekim", "nakit kullanım"]):
        scopes.add("credit_card_cash_advance")

    # KMH / Destek Hesap
    if any(k in name for k in ["destek hesap", "kmh", "ek hesap", "esnek hesap"]):
        scopes.add("overdraft")

    # Ücret / tarife
    if any(k in name for k in ["ücret", "tarife", "komisyon", "masraf"]):
        scopes.add("whitelist_and_pricing")

    # Elektronik / dijital
    if any(k in name for k in ["elektronik", "internet", "mobil", "dijital", "online", "fatura", "bilgilendirme", "aydınlatma"]):
        scopes.add("disclosure_and_documentation")

    # ATM / kasa
    if "atm" in name:
        scopes.add("atm")
    if "kasa" in name:
        scopes.add("safe_deposit_box")
    
    # Cayma / tüketici koruma
    if any(k in name for k in ["cayma", "tüketici", "sözleşme öncesi"]):
        scopes.add("disclosure_and_documentation")
    
    # Temerrüt / gecikme
    if any(k in name for k in ["temerrüt", "gecikme", "ödeme"]) and ("kart" in name or doc_id and "kart" in doc_id.lower()):
        scopes.add("credit_card")

    return sorted(scopes) if scopes else ["general_contract"]

# =============================================================================
# WORD PARSER
# =============================================================================

def parse_word_document(word_path: Path, doc_id: str) -> List[Dict[str, Any]]:
    """Parse Word document with structural awareness using styles."""
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")

    logger.info(f"Parsing Word: {word_path}")

    doc = Document(str(word_path))
    units = []

    current_section_name = None
    current_article_no = None
    current_article_title = None
    current_lines = []
    unit_counter = 1
    para_start_idx = 0

    last_text = None
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Shadow text/duplicate protection
        if text == last_text:
            continue
        last_text = text

        style_name = paragraph.style.name if paragraph.style else "Normal"

        # Heading 1 = new section
        if "Heading 1" in style_name or "heading 1" in style_name.lower(): # pyright: ignore[reportOptionalMemberAccess, reportOperatorIssue]
            # Save previous unit
            if current_lines:
                full_text = clean_noise_tokens(" ".join(current_lines))
                if full_text and len(full_text) > 10:
                    units.append({
                        "doc_id": doc_id,
                        "unit_id": f"{doc_id}_unit_{unit_counter}",
                        "unit_type": "section" if current_section_name else "paragraph",
                        "doc_type": determine_doc_type(doc_id),
                        "section_name": current_section_name,
                        "article_no": current_article_no,
                        "article_title": current_article_title,
                        "text": full_text,
                        "para_start": para_start_idx,
                        "para_end": para_idx - 1,
                        "metadata": {"style": style_name}
                    })
                    unit_counter += 1
                current_lines = []

            current_section_name = text
            current_article_no = None
            current_article_title = None
            para_start_idx = para_idx
            continue

        if madde_match := MADDE_RE.match(text):
            # Save previous unit
            if current_lines:
                full_text = clean_noise_tokens(" ".join(current_lines))
                if full_text and len(full_text) > 10:
                    units.append({
                        "doc_id": doc_id,
                        "unit_id": f"{doc_id}_unit_{unit_counter}",
                        "unit_type": "article",
                        "doc_type": determine_doc_type(doc_id),
                        "section_name": current_section_name,
                        "article_no": current_article_no,
                        "article_title": current_article_title,
                        "text": full_text,
                        "para_start": para_start_idx,
                        "para_end": para_idx - 1,
                        "metadata": {"style": style_name}
                    })
                    unit_counter += 1
                current_lines = []

            current_article_no = int(madde_match.group(1))
            current_article_title = madde_match.group(2).strip() if madde_match.group(2) else None
            para_start_idx = para_idx
            continue

        # List Paragraph = list item başlangıcı, devamındaki Body Text'leri de topla
        if "List" in style_name: # pyright: ignore[reportOperatorIssue]
            # Save previous unit if it exists
            if current_lines:
                full_text = clean_noise_tokens(" ".join(current_lines))
                if full_text and len(full_text) > 10:
                    units.append({
                        "doc_id": doc_id,
                        "unit_id": f"{doc_id}_unit_{unit_counter}",
                        "unit_type": "list_item",
                        "doc_type": determine_doc_type(doc_id),
                        "section_name": current_section_name,
                        "article_no": current_article_no,
                        "article_title": current_article_title,
                        "text": full_text,
                        "para_start": para_start_idx,
                        "para_end": para_idx - 1,
                        "metadata": {"style": "List Paragraph"}
                    })
                    unit_counter += 1
                current_lines = []

            # Start new list item, accumulate its text
            current_lines = [text]
            para_start_idx = para_idx
            continue

        # Regular content - accumulate
        current_lines.append(text)

    # Save final unit
    if current_lines:
        full_text = clean_noise_tokens(" ".join(current_lines))
        if full_text and len(full_text) > 10:
            units.append({
                "doc_id": doc_id,
                "unit_id": f"{doc_id}_unit_{unit_counter}",
                "unit_type": "paragraph",
                "doc_type": determine_doc_type(doc_id),
                "section_name": current_section_name,
                "article_no": current_article_no,
                "article_title": current_article_title,
                "text": full_text,
                "para_start": para_start_idx,
                "para_end": len(doc.paragraphs) - 1,
                "metadata": {"style": "Body Text"}
            })

    logger.info(f"Extracted {len(units)} units from Word")
    return units

# =============================================================================
# CHUNKING FUNCTIONS
# =============================================================================

def create_chunks(parsed_units: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Create chunks from parsed units."""
    chunks = []
    
    for unit in parsed_units:
        unit_id = unit.get("unit_id")
        if not unit_id:
            continue
        
        text = unit.get("text", "")
        if not text or len(text) < 10:
            continue

        scope = assign_bank_document_scopes(unit.get("doc_id", ""), text)
        
        chunk = {
            "chunk_id": unit_id,
            "doc_id": unit.get("doc_id"),
            "unit_id": unit_id,
            "unit_type": unit.get("unit_type"),
            "doc_type": unit.get("doc_type"),
            "section_name": unit.get("section_name"),
            "article_no": unit.get("article_no"),
            "article_title": unit.get("article_title"),
            "para_start": unit.get("para_start"),
            "para_end": unit.get("para_end"),
            "text": text,
            "scope": scope,
            "metadata": unit.get("metadata", {})
        }
        chunks.append(chunk)
    
    logger.info(f"Created {len(chunks)} chunks")
    return chunks

def create_vector_records(chunk: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Create vector-ready records from a chunk. Returns ONLY child chunks."""
    full_text = chunk.get("text", "")
    if not full_text:
        return []
    
    parent_id = chunk["chunk_id"]
    records = []
    
    # If text is short enough, create single child
    if len(full_text) <= MAX_EMBEDDING_CHARS:
        child = dict(chunk)
        child["chunk_id"] = f"{parent_id}__c1"
        child["chunk_kind"] = "child"
        child["parent_chunk_id"] = parent_id
        child["text_for_embedding"] = full_text
        child["text"] = full_text  # Fix: Ensure text matches embedding text
        records.append(child)
        return records
    
    # Split long text by sentences
    sentences = re.split(r'(?<=[\.!\?])\s+', full_text)
    
    current_chunk = ""
    child_idx = 1
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 <= MAX_EMBEDDING_CHARS:
            current_chunk = f"{current_chunk} {sentence}".strip() if current_chunk else sentence
        else:
            if current_chunk:
                child = dict(chunk)
                child["chunk_id"] = f"{parent_id}__c{child_idx}"
                child["chunk_kind"] = "child"
                child["parent_chunk_id"] = parent_id
                child["text_for_embedding"] = current_chunk
                child["text"] = current_chunk  # Fix: Overwrite parent text
                records.append(child)
                child_idx += 1
            current_chunk = sentence
    
    # Add remaining chunk
    if current_chunk:
        child = dict(chunk)
        child["chunk_id"] = f"{parent_id}__c{child_idx}"
        child["chunk_kind"] = "child"
        child["parent_chunk_id"] = parent_id
        child["text_for_embedding"] = current_chunk
        child["text"] = current_chunk  # Fix: Overwrite parent text
        records.append(child)
    
    return records

def perform_chunk_qc(chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Perform quality control on chunks."""
    chunk_ids = [chunk.get("chunk_id") for chunk in chunks if chunk.get("chunk_id")]
    duplicate_ids = [k for k, v in Counter(chunk_ids).items() if v > 1]
    
    empty_texts = [chunk.get("chunk_id") for chunk in chunks if not (chunk.get("text") or "").strip()]
    
    text_lengths = [(chunk.get("chunk_id"), len(chunk.get("text", ""))) for chunk in chunks]
    sorted_lengths = sorted(text_lengths, key=lambda x: x[1], reverse=True)
    
    unit_types = Counter(chunk.get("unit_type") for chunk in chunks)
    
    return {
        "counts": {
            "chunks": len(chunks),
            "unit_types": dict(unit_types),
        },
        "duplicate_chunk_ids": duplicate_ids,
        "empty_text": empty_texts,
        "text_length_stats": {
            "min": min((length for _, length in sorted_lengths), default=0),
            "max": max((length for _, length in sorted_lengths), default=0),
            "top_10_longest": sorted_lengths[:10],
        },
    }

def perform_vector_qc(records: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Perform quality control on vector records."""
    chunk_ids = [record.get("chunk_id") for record in records if record.get("chunk_id")]
    duplicate_ids = [k for k, v in Counter(chunk_ids).items() if v > 1]
    
    children = [record for record in records if record.get("chunk_kind") == "child"]
    too_long = [record["chunk_id"] for record in children
               if len(record.get("text_for_embedding", "")) > MAX_EMBEDDING_CHARS]
    too_short = [record["chunk_id"] for record in children
                if 0 < len(record.get("text_for_embedding", "")) < MIN_SUBCHUNK_CHARS]
    missing_parent = [record["chunk_id"] for record in children
                     if not record.get("parent_chunk_id")]
    
    parents = {record["chunk_id"] for record in records if record.get("chunk_kind") == "parent"}
    orphan_children = [record["chunk_id"] for record in children
                      if record.get("parent_chunk_id") not in parents]
    
    return {
        "counts": {
            "records_total": len(records),
            "parents": len(parents),
            "children": len(children),
            "duplicate_chunk_ids": len(duplicate_ids),
            "child_too_long": len(too_long),
            "child_too_short": len(too_short),
            "child_missing_parent_id": len(missing_parent),
            "child_orphan": len(orphan_children),
        },
        "sample": {
            "duplicate_chunk_ids": duplicate_ids[:50],
            "child_too_long": too_long[:50],
            "child_too_short": too_short[:50],
            "child_missing_parent_id": missing_parent[:50],
            "child_orphan": orphan_children[:50],
        },
    }

# =============================================================================
# FILE I/O FUNCTIONS
# =============================================================================

def save_json(data: Any, file_path: Path) -> None:
    """Save data to JSON file."""
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"Saved JSON to {file_path}")
    except Exception as e:
        logger.error(f"Error saving JSON to {file_path}: {e}")
        raise

def save_jsonl(data: List[Dict[str, Any]], file_path: Path) -> None:
    """Save list of dictionaries to JSONL file."""
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            for item in data:
                f.write(json.dumps(item, ensure_ascii=False) + "\n")
        logger.info(f"Saved JSONL to {file_path}")
    except Exception as e:
        logger.error(f"Error saving JSONL to {file_path}: {e}")
        raise

# =============================================================================
# MAIN EXECUTION FUNCTIONS
# =============================================================================

def process_word_document(file_path: Path) -> None:
    """Process a single Word document."""
    logger.info(f"Processing Word document: {file_path.name}")

    # Generate doc_id from filename
    doc_id = file_path.stem

    # Parse Word
    units = parse_word_document(file_path, doc_id)

    # Save parsed units
    parsed_output = BANKA_DOKUMANLARI_PARSED_DIR / f"{doc_id}_units.json" # pyright: ignore[reportOptionalOperand]
    save_json(units, parsed_output)

    # Create chunks
    chunks = create_chunks(units)
    _extracted_from_process_word_document_17(
        doc_id, '_chunks.json', '_chunks.jsonl', chunks
    )
    # QC on chunks
    chunk_qc = perform_chunk_qc(chunks)
    chunk_qc_output = BANKA_DOKUMANLARI_CHUNKS_DIR / f"{doc_id}_chunks_qc.json" # pyright: ignore[reportOptionalOperand]
    save_json(chunk_qc, chunk_qc_output)

    # Create vector-ready records
    records = []
    for chunk in chunks:
        records.extend(create_vector_records(chunk))

    # Attach excel vector records when the Word references an excel filename
    word_full_text = " ".join(normalize_for_embedding(c.get("text", "")) for c in chunks)
    try:
        base_dir = find_project_root(file_path.parent)
    except Exception:
        base_dir = find_project_root()
    if excel_records := find_referenced_excel_records(
        word_full_text, base_dir, doc_id
    ):
        records.extend(excel_records)
        logger.info(f"Attached {len(excel_records)} excel vector records for {doc_id}")

    _extracted_from_process_word_document_17(
        doc_id, '_vector_ready.json', '_vector_ready.jsonl', records
    )
    # QC on vector records
    vector_qc = perform_vector_qc(records)
    vector_qc_output = BANKA_DOKUMANLARI_CHUNKS_DIR / f"{doc_id}_vector_ready_qc.json" # pyright: ignore[reportOptionalOperand]
    save_json(vector_qc, vector_qc_output)

    logger.info(f"Completed processing: {file_path.name}")
    logger.info(f"  - Units: {len(units)}")
    logger.info(f"  - Chunks: {len(chunks)}")
    logger.info(f"  - Vector records: {len(records)}")
    logger.info(f"  - QC Summary: {chunk_qc['counts']}")


# TODO Rename this here and in `process_word_document`
def _extracted_from_process_word_document_17(doc_id, arg1, arg2, arg3):
    chunks_output = BANKA_DOKUMANLARI_CHUNKS_DIR / f"{doc_id}{arg1}" # pyright: ignore[reportOptionalOperand]
    chunks_jsonl_output = BANKA_DOKUMANLARI_CHUNKS_DIR / f"{doc_id}{arg2}" # pyright: ignore[reportOptionalOperand]
    save_json(arg3, chunks_output)
    save_jsonl(arg3, chunks_jsonl_output)

def main():
    """Main execution function."""
    try:
        # Setup
        base_dir = find_project_root()
        setup_directories(base_dir)
        log_file = base_dir / "logs" / "banka_dokumanlari_preparation_word.log"
        global logger
        logger = setup_logging(log_file)
        
        logger.info("Starting Banka Dokümanları Preparation - Word Only")
        logger.info(f"Base directory: {base_dir}")
        logger.info(f"Raw directory: {BANKA_DOKUMANLARI_RAW_DIR}")
        
        # Process all Word documents in raw directory
        if not BANKA_DOKUMANLARI_RAW_DIR.exists(): # pyright: ignore[reportOptionalMemberAccess]
            logger.error(f"Raw directory does not exist: {BANKA_DOKUMANLARI_RAW_DIR}")
            return
        
        # Find all Word documents
        word_documents = [f for f in BANKA_DOKUMANLARI_RAW_DIR.iterdir() # pyright: ignore[reportOptionalMemberAccess] 
                        if f.is_file() and f.suffix.lower() == '.docx']
        
        if not word_documents:
            logger.warning("No Word documents found in raw directory")
            return
        
        logger.info(f"Found {len(word_documents)} Word documents to process")
        
        for doc_path in word_documents:
            try:
                process_word_document(doc_path)
            except Exception as e:
                logger.error(f"Failed to process {doc_path.name}: {e}")
                continue
        
        logger.info("All Word documents processed successfully")
        
    except Exception as e:
        logger.error(f"Error in main execution: {e}")
        raise

if __name__ == "__main__":
    main()
